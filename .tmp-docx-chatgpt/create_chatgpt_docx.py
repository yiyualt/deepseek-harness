from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("/Users/yuyi/Desktop/chatgpt.docx")
TITLE = "雨后的一天"
DATE = "2026年8月17日　星期一　天气：雨后晴"
PARAGRAPHS = [
    "清晨推开窗，昨夜的雨刚停，树叶上挂着亮晶晶的水珠，空气里有泥土和青草的清香。我背起书包出门，故意绕到河边，看白鹭从浅水中轻轻飞起。",
    "上学路上，我遇见一位推车上坡的老人，便跑过去帮忙。老人笑着道谢，我心里也暖暖的。下午放学，夕阳把积水照成金色，几个孩子踩着水洼欢笑。",
    "我没有急着回家，而是慢慢走，听风吹过梧桐叶的沙沙声。回到家，我把今天的小事写进日记。睡前想起这些画面仍觉得满足。",
    "原来平凡的一天，只要认真观察、愿意伸手帮助别人，也会藏着许多值得珍惜的光亮。",
]


def set_east_asia_font(run, name: str) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)


body = "".join(PARAGRAPHS)
assert len(re.findall(r"[\u4e00-\u9fff]", body)) == 200

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Arial Unicode MS"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
normal.font.size = Pt(12)
normal.font.color.rgb = RGBColor(45, 52, 54)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(18)
title.paragraph_format.space_after = Pt(6)
title.paragraph_format.keep_with_next = True
title_run = title.add_run(TITLE)
set_east_asia_font(title_run, "Arial Unicode MS")
title_run.font.size = Pt(22)
title_run.bold = True
title_run.font.color.rgb = RGBColor(47, 93, 80)

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
date.paragraph_format.space_before = Pt(0)
date.paragraph_format.space_after = Pt(22)
date.paragraph_format.keep_with_next = True
date_run = date.add_run(DATE)
set_east_asia_font(date_run, "Arial Unicode MS")
date_run.font.size = Pt(10.5)
date_run.font.color.rgb = RGBColor(102, 112, 108)

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(0)
rule.paragraph_format.space_after = Pt(18)
p_pr = rule._p.get_or_add_pPr()
p_bdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "8")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "A9C7B8")
p_bdr.append(bottom)
p_pr.append(p_bdr)

for text in PARAGRAPHS:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.333
    run = paragraph.add_run(text)
    set_east_asia_font(run, "Arial Unicode MS")
    run.font.size = Pt(12)

doc.core_properties.title = TITLE
doc.core_properties.subject = "约200字中文日记"
doc.core_properties.author = ""
doc.core_properties.keywords = "日记"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(f"created={OUTPUT}")
print("chinese_characters=200")
